"""Build the synthetic 40-document extraction golden corpus."""

from __future__ import annotations

import hashlib
import io
import json
import shutil
from pathlib import Path


ROOT = Path(__file__).resolve().parent
GENERATED = ROOT / "generated"
MANIFEST = json.loads((ROOT / "manifest.json").read_text(encoding="utf-8"))


def _font(size: int):
    from PIL import ImageFont

    candidates = [
        Path("C:/Windows/Fonts/Nirmala.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _scan_image(index: int, *, mixed: bool = False):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(42)
    body_font = _font(30)
    lines = [
        "FIRST INFORMATION REPORT",
        f"FIR No. {120 + index}/2024",
        "Police Station: Nyay Nagar",
        "Complainant: Synthetic Corpus Record",
        "Section 154 CrPC - registration requested",
        "Date of occurrence: 12/03/2024",
        "This fixture contains no real personal information.",
    ]
    if mixed:
        lines = [
            "NyaySetu bilingual legal record",
            "Section 154 CrPC - FIR registration",
            "Nyaya aur kanooni sahayata / legal assistance",
            "Hindi-English mixed document fixture",
            "\u0928\u094d\u092f\u093e\u092f \u0938\u0939\u093e\u092f\u0924\u093e \u0914\u0930 \u0915\u093e\u0928\u0942\u0928\u0940 \u0905\u092d\u093f\u0932\u0947\u0916",
        ]
    y = 110
    for line_index, line in enumerate(lines):
        draw.text(
            (90, y),
            line,
            fill="black",
            font=title_font if line_index == 0 else body_font,
        )
        y += 105 if line_index == 0 else 82
    draw.rectangle((70, 70, 1170, 1660), outline="black", width=4)
    draw.line((80, 760, 1160, 760), fill="black", width=3)
    if index % 3 == 0:
        image = image.rotate(2.0, expand=True, fillcolor="white")
    if index % 8 == 0:
        image = image.rotate(90, expand=True, fillcolor="white")
    return image


def _save_scanned_pdf(path: Path, index: int):
    import fitz

    image = _scan_image(index)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    image.close()

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_image(page.rect, stream=buffer.getvalue())
    document.save(path, deflate=True)
    document.close()


def _save_judgment(path: Path, index: int):
    import fitz

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    text = (
        "IN THE HIGH COURT OF DELHI\n\n"
        f"WRIT PETITION (C) NO. {1000 + index} OF 2024\n\n"
        "The petition challenges an administrative order. The Court heard "
        "the petitioner and the respondent, examined the complete record, "
        "and recorded reasons consistent with natural justice.\n\n"
        "ORDER\nThe petition is disposed of with directions to the authority "
        "to decide the representation by a reasoned order within four weeks."
    )
    page.insert_textbox(fitz.Rect(55, 65, 540, 780), text, fontsize=11)
    document.set_metadata(
        {
            "title": f"Synthetic judgment {index}",
            "author": "NyaySetu golden corpus",
        }
    )
    document.save(path, deflate=True)
    document.close()


def _save_form(path: Path, index: int):
    import fitz

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_text((55, 65), "LEGAL AID INTAKE FORM", fontsize=16)
    labels = [
        ("Applicant", "Synthetic Applicant"),
        ("Case Number", f"LA-{index:03d}/2024"),
        ("District", "New Delhi"),
        ("Relief Requested", "Registration and certified copy"),
        ("Applicable Provision", "Section 154 CrPC"),
    ]
    y = 120
    for label, value in labels:
        page.insert_text((60, y), f"{label}:", fontsize=10)
        page.draw_line((170, y + 2), (520, y + 2), width=0.8)
        page.insert_text((180, y), value, fontsize=10)
        y += 58
    page.draw_rect(fitz.Rect(55, 95, 535, y), width=1)
    document.save(path, deflate=True)
    document.close()


def _save_table_pdf(path: Path, index: int):
    import fitz

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_text((55, 55), "EVIDENCE AND PROVISION TABLE", fontsize=15)
    x_positions = [55, 210, 390, 540]
    y_positions = [90, 130, 170, 210, 250]
    for x in x_positions:
        page.draw_line((x, y_positions[0]), (x, y_positions[-1]), width=1)
    for y in y_positions:
        page.draw_line((x_positions[0], y), (x_positions[-1], y), width=1)
    rows = [
        ["Provision", "Evidence", "Finding"],
        ["154 CrPC", "Complaint", "FIR required"],
        ["156(3) CrPC", "Affidavit", "Magistrate review"],
        ["Article 14", f"Record {index}", "Equal treatment"],
    ]
    for row_index, row in enumerate(rows):
        y = y_positions[row_index] + 25
        for column_index, value in enumerate(row):
            page.insert_text(
                (x_positions[column_index] + 6, y),
                value,
                fontsize=8,
            )
    document.save(path, deflate=True)
    document.close()


def _save_mixed_pdf(path: Path, index: int):
    import fitz

    document = fitz.open()
    digital = document.new_page(width=595, height=842)
    text = (
        "NyaySetu bilingual case summary\n\n"
        "This digital page records a request under Section 154 CrPC. "
        "The following scanned page contains Hindi and English text for OCR."
    )
    digital.insert_textbox(fitz.Rect(55, 70, 540, 760), text, fontsize=12)

    image = _scan_image(index, mixed=True)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    image.close()
    scanned = document.new_page(width=595, height=842)
    scanned.insert_image(scanned.rect, stream=buffer.getvalue())
    document.save(path, deflate=True)
    document.close()


def _save_docx(path: Path, index: int):
    from docx import Document

    document = Document()
    document.add_heading("LEGAL MEMORANDUM", level=1)
    document.add_paragraph(
        "The memorandum reviews a request under Section 154 CrPC and "
        "preserves every proposition with a source reference."
    )
    table = document.add_table(rows=3, cols=2)
    values = [
        ("Provision", "Finding"),
        ("Section 154 CrPC", "Registration required"),
        ("Record", f"Memo {index}"),
    ]
    for row, values_row in zip(table.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value
    document.save(path)


def _save_xlsx(path: Path, index: int):
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Case Register"
    rows = [
        ["Case Number", "Status", "Provision"],
        [f"CASE-{index:03d}", "Open", "Section 154 CrPC"],
        [f"CASE-{index + 100:03d}", "Reviewed", "Article 14"],
    ]
    for row in rows:
        sheet.append(row)
    workbook.save(path)
    workbook.close()


def _save_image(path: Path, index: int):
    image = _scan_image(index)
    image.save(path, format="PNG", optimize=True)
    image.close()


def build() -> None:
    if GENERATED.exists():
        shutil.rmtree(GENERATED)
    GENERATED.mkdir(parents=True)
    checksums = {}
    category_counts: dict[str, int] = {}

    for item in MANIFEST["documents"]:
        category = item["category"]
        category_counts[category] = category_counts.get(category, 0) + 1
        index = category_counts[category]
        path = GENERATED / item["filename"]
        if category == "scanned_fir" and item["format"] == "pdf":
            _save_scanned_pdf(path, index)
        elif category == "judgment":
            _save_judgment(path, index)
        elif category == "form":
            _save_form(path, index)
        elif category == "table":
            _save_table_pdf(path, index)
        elif category == "mixed_language":
            _save_mixed_pdf(path, index)
        elif category == "docx":
            _save_docx(path, index)
        elif category == "spreadsheet":
            _save_xlsx(path, index)
        elif item["format"] == "image":
            _save_image(path, index)
        else:
            raise ValueError(f"Unsupported corpus item: {item['id']}")
        checksums[item["filename"]] = hashlib.sha256(path.read_bytes()).hexdigest()

    (ROOT / "checksums.json").write_text(
        json.dumps(checksums, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    print(f"Built {len(checksums)} golden documents in {GENERATED}")


if __name__ == "__main__":
    build()
