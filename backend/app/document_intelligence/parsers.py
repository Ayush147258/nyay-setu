"""Multi-format parser router producing the canonical Document IR blocks."""

from __future__ import annotations

import csv
import hashlib
import io
import logging
import mimetypes
import re
import zipfile
from dataclasses import dataclass, field
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any

from app.document_intelligence.pdf_pipeline import normalize_block_text
from app.document_intelligence.models import (
    BlockKind,
    DocumentBlock,
    DocumentFormat,
    DocumentPage,
    ParseStatus,
    StructuredTable,
    TableCell,
    TableRow,
)


logger = logging.getLogger(__name__)
PARSER_VERSION = "2.0.0"
_MAX_ZIP_ENTRIES = 10_000
_MAX_ZIP_UNCOMPRESSED = 250 * 1024 * 1024


class ParseError(RuntimeError):
    def __init__(
        self,
        message: str,
        *,
        status: ParseStatus = ParseStatus.FAILED,
        error_code: str = "parse_error",
        metadata: dict[str, Any] | None = None,
    ):
        super().__init__(message)
        self.status = status
        self.error_code = error_code
        self.metadata = metadata or {}


@dataclass
class ParseOutput:
    document_format: DocumentFormat
    status: ParseStatus
    parser_name: str
    pages: list[DocumentPage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


def _block_id(version_id: str, page: int, sequence: int, text: str) -> str:
    value = f"{version_id}:{page}:{sequence}:{text}".encode("utf-8", errors="ignore")
    return "blk_" + hashlib.sha256(value).hexdigest()[:20]


def _blocks_from_text(
    text: str,
    *,
    version_id: str,
    page_number: int,
    confidence: float = 1.0,
    kind: BlockKind = BlockKind.PARAGRAPH,
) -> list[DocumentBlock]:
    normalized = normalize_block_text(text)
    chunks = [
        normalize_block_text(chunk)
        for chunk in re.split(r"\n\s*\n|(?<=\.)\s*\n", normalized)
        if normalize_block_text(chunk)
    ]
    return [
        DocumentBlock(
            block_id=_block_id(version_id, page_number, sequence, chunk),
            page_number=page_number,
            sequence=sequence,
            kind=kind,
            text=chunk,
            confidence=confidence,
            metadata={
                "char_length": len(chunk),
                "source_span_validated": True,
            },
        )
        for sequence, chunk in enumerate(chunks)
    ]


def _table_block_from_rows(
    rows: list[list[Any]],
    *,
    version_id: str,
    page_number: int,
    sequence: int = 0,
    confidence: float = 1.0,
) -> DocumentBlock | None:
    normalized = [
        [
            re.sub(r"\s+", " ", "" if cell is None else str(cell)).strip()
            for cell in row
        ]
        for row in rows
    ]
    if not any(any(cell for cell in row) for row in normalized):
        return None
    text = "\n".join(" | ".join(row).rstrip() for row in normalized).strip()
    table_id = (
        "tbl_"
        + hashlib.sha256(
            f"{version_id}:{page_number}:{sequence}:{text}".encode(
                "utf-8",
                errors="ignore",
            )
        ).hexdigest()[:20]
    )
    structured = StructuredTable(
        table_id=table_id,
        page_number=page_number,
        rows=[
            TableRow(
                row_index=row_index,
                cells=[
                    TableCell(
                        row_index=row_index,
                        column_index=column_index,
                        text=cell,
                    )
                    for column_index, cell in enumerate(row)
                ],
            )
            for row_index, row in enumerate(normalized)
        ],
    )
    return DocumentBlock(
        block_id=_block_id(version_id, page_number, sequence, text),
        page_number=page_number,
        sequence=sequence,
        kind=BlockKind.TABLE,
        text=text,
        confidence=confidence,
        table=structured,
        metadata={
            "rows": len(structured.rows),
            "columns": max(
                (len(row.cells) for row in structured.rows),
                default=0,
            ),
            "source_span_validated": True,
        },
    )


def _inspect_zip(path: Path) -> set[str]:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > _MAX_ZIP_ENTRIES:
                raise ParseError("Archive contains too many entries")
            total = sum(entry.file_size for entry in entries)
            if total > _MAX_ZIP_UNCOMPRESSED:
                raise ParseError("Archive expands beyond the safety limit")
            return {entry.filename for entry in entries}
    except zipfile.BadZipFile as exc:
        raise ParseError("Invalid ZIP-based document") from exc


def detect_format(path: Path, filename: str, declared_media_type: str) -> tuple[DocumentFormat, str]:
    head = path.read_bytes()[:32]
    suffix = Path(filename).suffix.casefold()
    media_type = (declared_media_type or mimetypes.guess_type(filename)[0] or "").casefold()

    if head.startswith(b"%PDF-"):
        return DocumentFormat.PDF, "application/pdf"
    if head.startswith(b"PK\x03\x04"):
        names = _inspect_zip(path)
        if any(name.startswith("word/") for name in names):
            return DocumentFormat.DOCX, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if any(name.startswith("xl/") for name in names):
            return DocumentFormat.SPREADSHEET, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        return DocumentFormat.UNKNOWN, media_type or "application/zip"
    if head.startswith(b"\x89PNG\r\n\x1a\n") or head.startswith(b"\xff\xd8\xff") or head[:4] in {
        b"II*\x00",
        b"MM\x00*",
    }:
        return DocumentFormat.IMAGE, media_type or "image/*"
    if suffix == ".csv" or media_type == "text/csv":
        return DocumentFormat.CSV, "text/csv"
    if suffix in {".txt", ".md", ".rtf"} or media_type.startswith("text/plain"):
        return DocumentFormat.TEXT, media_type or "text/plain"
    if suffix == ".eml" or media_type == "message/rfc822":
        return DocumentFormat.EMAIL, "message/rfc822"
    if suffix in {".xls", ".xlsx"}:
        return DocumentFormat.SPREADSHEET, media_type or "application/vnd.ms-excel"
    if suffix in {".jpg", ".jpeg", ".png", ".tif", ".tiff"} or media_type.startswith("image/"):
        return DocumentFormat.IMAGE, media_type
    return DocumentFormat.UNKNOWN, media_type or "application/octet-stream"


class ParserRouter:
    def __init__(self, pdf_limits: Any | None = None):
        self.pdf_limits = pdf_limits

    def parse(self, path: Path, filename: str, media_type: str, version_id: str) -> ParseOutput:
        document_format, detected_media_type = detect_format(path, filename, media_type)
        try:
            if document_format == DocumentFormat.PDF:
                result = self._parse_pdf(path, version_id)
            elif document_format == DocumentFormat.DOCX:
                result = self._parse_docx(path, version_id)
            elif document_format == DocumentFormat.TEXT:
                result = self._parse_text(path, version_id)
            elif document_format == DocumentFormat.CSV:
                result = self._parse_csv(path, version_id)
            elif document_format == DocumentFormat.SPREADSHEET:
                result = self._parse_spreadsheet(path, filename, version_id)
            elif document_format == DocumentFormat.EMAIL:
                result = self._parse_email(path, version_id)
            elif document_format == DocumentFormat.IMAGE:
                result = self._parse_image(path, version_id)
            else:
                result = ParseOutput(
                    document_format=DocumentFormat.UNKNOWN,
                    status=ParseStatus.UNSUPPORTED,
                    parser_name="unsupported",
                    warnings=[
                        "Unsupported format. Preserve the original and convert it to PDF, DOCX, text, image, spreadsheet, or EML."
                    ],
                )
            result.metadata["detected_media_type"] = detected_media_type
            return result
        except ParseError as exc:
            return ParseOutput(
                document_format=document_format,
                status=exc.status,
                parser_name="parser_router",
                warnings=[str(exc)],
                metadata={
                    "detected_media_type": detected_media_type,
                    "error_code": exc.error_code,
                    **exc.metadata,
                },
            )
        except Exception as exc:
            logger.exception(
                "Unexpected parser failure | format=%s",
                document_format.value,
            )
            return ParseOutput(
                document_format=document_format,
                status=ParseStatus.FAILED,
                parser_name="parser_router",
                warnings=[f"Parser failed safely: {type(exc).__name__}"],
                metadata={"detected_media_type": detected_media_type},
            )

    def _parse_pdf(self, path: Path, version_id: str) -> ParseOutput:
        from app.document_intelligence.pdf_pipeline import (
            PdfExtractionError,
            PdfExtractionPipeline,
        )

        try:
            result = PdfExtractionPipeline(self.pdf_limits).extract(
                path,
                version_id,
            )
        except PdfExtractionError as exc:
            raise ParseError(
                str(exc),
                status=exc.status,
                error_code=exc.error_code,
                metadata=exc.metadata,
            ) from exc
        return ParseOutput(
            DocumentFormat.PDF,
            result.status,
            result.parser_name,
            result.pages,
            result.warnings,
            result.metadata,
        )
    def _parse_docx(self, path: Path, version_id: str) -> ParseOutput:
        try:
            from docx import Document
        except ImportError as exc:
            raise ParseError("DOCX support is not installed") from exc
        try:
            document = Document(str(path))
            blocks: list[DocumentBlock] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                style = (paragraph.style.name if paragraph.style else "").casefold()
                kind = BlockKind.HEADING if "heading" in style or "title" in style else BlockKind.PARAGRAPH
                sequence = len(blocks)
                blocks.append(
                    DocumentBlock(
                        block_id=_block_id(version_id, 1, sequence, text),
                        page_number=1,
                        sequence=sequence,
                        kind=kind,
                        text=text,
                    )
                )
            for table in document.tables:
                table_block = _table_block_from_rows(
                    [
                        [cell.text for cell in row.cells]
                        for row in table.rows
                    ],
                    version_id=version_id,
                    page_number=1,
                    sequence=len(blocks),
                )
                if table_block is not None:
                    blocks.append(table_block)
            return ParseOutput(
                DocumentFormat.DOCX,
                ParseStatus.READY if blocks else ParseStatus.PARTIAL,
                "python-docx",
                [DocumentPage(
                    page_number=1,
                    confidence=1.0 if blocks else 0.0,
                    extraction_method="digital" if blocks else "none",
                    blocks=blocks,
                )],
                [] if blocks else ["DOCX contained no extractable text"],
                {"page_boundaries": "not_available_in_docx"},
            )
        except Exception as exc:
            raise ParseError(f"DOCX could not be read: {exc}") from exc

    def _decode_text(self, path: Path) -> tuple[str, str]:
        raw = path.read_bytes()
        for encoding in ("utf-8-sig", "utf-16", "cp1252"):
            try:
                return raw.decode(encoding), encoding
            except UnicodeDecodeError:
                continue
        raise ParseError("Text encoding is unsupported")

    def _parse_text(self, path: Path, version_id: str) -> ParseOutput:
        text, encoding = self._decode_text(path)
        blocks = _blocks_from_text(text, version_id=version_id, page_number=1)
        return ParseOutput(
            DocumentFormat.TEXT,
            ParseStatus.READY if blocks else ParseStatus.PARTIAL,
            "text",
            [DocumentPage(
                    page_number=1,
                    confidence=1.0 if blocks else 0.0,
                    extraction_method="digital" if blocks else "none",
                    blocks=blocks,
                )],
            [] if blocks else ["Text document was empty"],
            {"encoding": encoding},
        )

    def _parse_csv(self, path: Path, version_id: str) -> ParseOutput:
        text, encoding = self._decode_text(path)
        rows = list(csv.reader(io.StringIO(text)))
        table = _table_block_from_rows(
            rows,
            version_id=version_id,
            page_number=1,
        )
        blocks = [table] if table is not None else []
        return ParseOutput(
            DocumentFormat.CSV,
            ParseStatus.READY if blocks else ParseStatus.PARTIAL,
            "csv",
            [DocumentPage(
                    page_number=1,
                    confidence=1.0 if blocks else 0.0,
                    extraction_method="digital" if blocks else "none",
                    blocks=blocks,
                )],
            [] if blocks else ["CSV document was empty"],
            {"encoding": encoding, "rows": len(rows)},
        )

    def _parse_spreadsheet(self, path: Path, filename: str, version_id: str) -> ParseOutput:
        suffix = Path(filename).suffix.casefold()
        if suffix == ".xls":
            return self._parse_xls(path, version_id)
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise ParseError("XLSX support is not installed") from exc
        try:
            workbook = load_workbook(path, read_only=True, data_only=True)
            pages: list[DocumentPage] = []
            sheet_names = [sheet.title for sheet in workbook.worksheets]
            for page_number, worksheet in enumerate(workbook.worksheets, start=1):
                rows = [
                    list(row)
                    for row in worksheet.iter_rows(values_only=True)
                ]
                table = _table_block_from_rows(
                    rows,
                    version_id=version_id,
                    page_number=page_number,
                )
                blocks = [table] if table is not None else []
                pages.append(
                    DocumentPage(
                        page_number=page_number,
                        confidence=1.0 if blocks else 0.0,
                        extraction_method="digital" if blocks else "none",
                        blocks=blocks,
                    )
                )
            workbook.close()
            has_blocks = any(page.blocks for page in pages)
            return ParseOutput(
                DocumentFormat.SPREADSHEET,
                ParseStatus.READY if has_blocks else ParseStatus.PARTIAL,
                "openpyxl",
                pages,
                [] if has_blocks else ["Spreadsheet contained no values"],
                {"sheets": sheet_names},
            )
        except Exception as exc:
            raise ParseError(f"XLSX could not be read: {exc}") from exc

    def _parse_xls(self, path: Path, version_id: str) -> ParseOutput:
        try:
            import xlrd
        except ImportError as exc:
            raise ParseError("Legacy XLS support is not installed") from exc
        try:
            workbook = xlrd.open_workbook(path)
            pages: list[DocumentPage] = []
            for page_number, worksheet in enumerate(workbook.sheets(), start=1):
                rows = [
                    [
                        worksheet.cell_value(row, column)
                        for column in range(worksheet.ncols)
                    ]
                    for row in range(worksheet.nrows)
                ]
                table = _table_block_from_rows(
                    rows,
                    version_id=version_id,
                    page_number=page_number,
                )
                blocks = [table] if table is not None else []
                pages.append(
                    DocumentPage(
                        page_number=page_number,
                        confidence=1.0 if blocks else 0.0,
                        extraction_method="digital" if blocks else "none",
                        blocks=blocks,
                    )
                )
            has_blocks = any(page.blocks for page in pages)
            return ParseOutput(
                DocumentFormat.SPREADSHEET,
                ParseStatus.READY if has_blocks else ParseStatus.PARTIAL,
                "xlrd",
                pages,
                [] if has_blocks else ["Spreadsheet contained no values"],
            )
        except Exception as exc:
            raise ParseError(f"XLS could not be read: {exc}") from exc

    def _parse_email(self, path: Path, version_id: str) -> ParseOutput:
        try:
            message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
            metadata_text = "\n".join(
                f"{name}: {message.get(name, '')}" for name in ("From", "To", "Cc", "Date", "Subject")
            )
            body = message.get_body(preferencelist=("plain",)) if message.is_multipart() else message
            body_text = body.get_content() if body else ""
            blocks = _blocks_from_text(
                metadata_text,
                version_id=version_id,
                page_number=1,
                kind=BlockKind.METADATA,
            )
            body_blocks = _blocks_from_text(body_text, version_id=version_id, page_number=1)
            for block in body_blocks:
                block.sequence += len(blocks)
                block.block_id = _block_id(version_id, 1, block.sequence, block.text)
            blocks.extend(body_blocks)
            return ParseOutput(
                DocumentFormat.EMAIL,
                ParseStatus.READY if blocks else ParseStatus.PARTIAL,
                "stdlib-email",
                [DocumentPage(
                    page_number=1,
                    confidence=1.0 if blocks else 0.0,
                    extraction_method="digital" if blocks else "none",
                    blocks=blocks,
                )],
                [] if blocks else ["Email contained no extractable content"],
                {"attachments": len(list(message.iter_attachments()))},
            )
        except Exception as exc:
            raise ParseError(f"Email could not be read: {exc}") from exc

    def _parse_image(self, path: Path, version_id: str) -> ParseOutput:
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return ParseOutput(
                DocumentFormat.IMAGE,
                ParseStatus.OCR_REQUIRED,
                "image-metadata",
                warnings=["OCR dependencies are unavailable on this worker"],
                metadata={"error_code": "ocr_dependency_missing"},
            )

        from app.document_intelligence.pdf_pipeline import (
            PdfExtractionError,
            PdfExtractionLimits,
            PdfExtractionPipeline,
        )

        limits = PdfExtractionLimits.from_settings()
        pipeline = PdfExtractionPipeline(limits)
        Image.MAX_IMAGE_PIXELS = limits.max_page_pixels
        pages: list[DocumentPage] = []
        warnings: list[str] = []
        total_pixels = 0
        try:
            image = Image.open(path)
        except Exception as exc:
            return ParseOutput(
                DocumentFormat.IMAGE,
                ParseStatus.CORRUPTED,
                "pillow",
                warnings=[f"Image is corrupted or unreadable: {type(exc).__name__}"],
                metadata={"error_code": "image_corrupted"},
            )

        try:
            frame_count = int(getattr(image, "n_frames", 1))
            if frame_count > limits.max_pages:
                return ParseOutput(
                    DocumentFormat.IMAGE,
                    ParseStatus.LIMIT_EXCEEDED,
                    "pillow",
                    warnings=[
                        f"Image contains {frame_count} frames; "
                        f"limit is {limits.max_pages}"
                    ],
                    metadata={
                        "error_code": "image_frame_limit",
                        "frame_count": frame_count,
                    },
                )

            for frame_index in range(frame_count):
                page_number = frame_index + 1
                try:
                    image.seek(frame_index)
                    frame_pixels = image.width * image.height
                    total_pixels += frame_pixels
                    if frame_pixels > limits.max_page_pixels:
                        raise PdfExtractionError(
                            f"Image frame {page_number} exceeds the pixel limit",
                            status=ParseStatus.LIMIT_EXCEEDED,
                            error_code="image_pixel_limit",
                        )
                    if total_pixels > limits.max_total_ocr_pixels:
                        raise PdfExtractionError(
                            "Image total pixel limit exceeded",
                            status=ParseStatus.LIMIT_EXCEEDED,
                            error_code="image_total_pixel_limit",
                        )
                    frame = image.convert("RGB")
                    prepared, orientation, deskew = pipeline.prepare_ocr_image(
                        frame,
                        pytesseract,
                    )
                    try:
                        data = pytesseract.image_to_data(
                            prepared,
                            lang=limits.ocr_languages,
                            config="--psm 3",
                            output_type=pytesseract.Output.DICT,
                        )
                        language_warning = ""
                    except Exception as primary_exc:
                        data = pytesseract.image_to_data(
                            prepared,
                            lang="eng",
                            config="--psm 3",
                            output_type=pytesseract.Output.DICT,
                        )
                        language_warning = (
                            "Configured OCR languages were unavailable; "
                            f"English fallback used after {type(primary_exc).__name__}"
                        )

                    words = [
                        normalize_block_text(str(word))
                        for word in data.get("text", [])
                        if normalize_block_text(str(word))
                    ]
                    confidences = [
                        float(value)
                        for value in data.get("conf", [])
                        if str(value).replace(".", "", 1).lstrip("-").isdigit()
                        and float(value) >= 0
                    ]
                    confidence = (
                        sum(confidences) / len(confidences) / 100
                        if confidences
                        else 0.0
                    )
                    blocks = _blocks_from_text(
                        " ".join(words),
                        version_id=version_id,
                        page_number=page_number,
                        confidence=confidence,
                        kind=BlockKind.IMAGE_TEXT,
                    )
                    page_warnings = []
                    if language_warning:
                        page_warnings.append(language_warning)
                    if not blocks:
                        page_warnings.append("OCR produced no usable text")
                    elif confidence < 0.65:
                        page_warnings.append(
                            f"OCR confidence {confidence:.2f} requires human review"
                        )
                    pages.append(
                        DocumentPage(
                            page_number=page_number,
                            width=prepared.width,
                            height=prepared.height,
                            rotation=orientation,
                            confidence=max(0.0, min(1.0, confidence)),
                            extraction_method="ocr" if blocks else "none",
                            warnings=page_warnings,
                            blocks=blocks,
                            metadata={
                                "orientation_correction": orientation,
                                "deskew_angle": deskew,
                                "pixels": frame_pixels,
                                "source_spans_validated": all(
                                    block.metadata.get("source_span_validated")
                                    for block in blocks
                                ),
                            },
                        )
                    )
                    warnings.extend(
                        f"Page {page_number}: {warning}"
                        for warning in page_warnings
                    )
                    if prepared is not frame:
                        prepared.close()
                    frame.close()
                except PdfExtractionError as exc:
                    return ParseOutput(
                        DocumentFormat.IMAGE,
                        exc.status,
                        "pillow+tesseract",
                        pages,
                        warnings + [str(exc)],
                        {
                            "error_code": exc.error_code,
                            "frame_count": frame_count,
                            "processed_frames": len(pages),
                        },
                    )
                except Exception as exc:
                    page_warning = (
                        f"OCR failed safely on frame {page_number}: "
                        f"{type(exc).__name__}"
                    )
                    warnings.append(page_warning)
                    pages.append(
                        DocumentPage(
                            page_number=page_number,
                            confidence=0.0,
                            extraction_method="none",
                            warnings=[page_warning],
                            metadata={"error_code": "image_ocr_failed"},
                        )
                    )

            usable = [page for page in pages if page.blocks]
            if not usable:
                status = ParseStatus.OCR_REQUIRED
            elif len(usable) != len(pages) or any(
                page.confidence < 0.65 for page in usable
            ):
                status = ParseStatus.PARTIAL
            else:
                status = ParseStatus.READY
            return ParseOutput(
                DocumentFormat.IMAGE,
                status,
                "pillow+tesseract",
                pages,
                warnings,
                {
                    "frame_count": frame_count,
                    "processed_incrementally": True,
                    "total_pixels": total_pixels,
                },
            )
        finally:
            image.close()